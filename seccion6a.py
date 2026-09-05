#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera seccion6a_v1.docx — Sección 6a: Diseño del Plan de Mejoras"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3)
sec.right_margin  = Cm(2)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)

for h_name, sz in [('Heading 1',14),('Heading 2',12),('Heading 3',11)]:
    s = doc.styles[h_name]
    s.font.name  = 'Arial'
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = RGBColor(0,0,0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def p(doc, text, bold=False, italic=False, size=11, space_after=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after  = Pt(space_after)
    par.paragraph_format.space_before = Pt(2)
    run = par.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    return par

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'AAAAAA')
        tcB.append(b)
    tcPr.append(tcB)

def header_row(table, headers, widths, bg='1F5C99'):
    row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, bg)
        set_borders(cell)
        par = cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(hdr)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size  = Pt(9)
        run.font.name  = 'Arial'

def data_row(table, values, widths, row_bg=None, sizes=None, bold_cols=None):
    row = table.add_row()
    for i, (val, w) in enumerate(zip(values, widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        if row_bg:
            set_cell_bg(cell, row_bg)
        set_borders(cell)
        par = cell.paragraphs[0]
        sz  = sizes[i] if sizes else 9
        run = par.add_run(str(val))
        run.font.size = Pt(sz)
        run.font.name = 'Arial'
        if bold_cols and i in bold_cols:
            run.font.bold = True

# ══════════════════════════════════════════════════════════════════════════════
# TÍTULO
# ══════════════════════════════════════════════════════════════════════════════
h(doc, '6. Desarrollo del Plan de Mejoras', level=1)
h(doc, '6a. Diseño del Plan de Mejoras Alineado a Estándares', level=2)

p(doc,
  'El presente plan de mejoras fue diseñado a partir de los hallazgos '
  'documentados en la Sección 5, producto de la ejecución de OWASP ZAP 2.17.0 '
  '(DAST) y SonarQube Community (SAST) sobre el sistema Eco-Alerta. El plan '
  'se estructura en conformidad con los siguientes estándares internacionales:',
  space_after=6)

# Tabla de estándares
ts = doc.add_table(rows=1, cols=3)
ts.style = 'Table Grid'
ts.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ts, ['Estándar','Nombre Completo','Aplicación en este Plan'], [2.5,5.0,8.0])
for vals in [
    ('ISO/IEC 25040','Systems and software engineering — Systems and software Quality Requirements '
     'and Evaluation (SQuaRE) — Evaluation process',
     'Marco para la evaluación sistemática de la calidad del producto software. '
     'Estructura la identificación, priorización y verificación de mejoras.'),
    ('ISO/IEC 12207','Systems and software engineering — Software life cycle processes',
     'Define los procesos del ciclo de vida del software. Orienta la implementación '
     'de mejoras dentro del proceso de desarrollo y mantenimiento.'),
    ('ISO/IEC 14764','Software Engineering — Software Life Cycle Processes — Maintenance',
     'Establece los procesos de mantenimiento del software. Guía la aplicación '
     'de mejoras correctivas, preventivas y perfectivas al sistema.'),
]:
    data_row(ts, vals, [2.5,5.0,8.0], sizes=[9,9,9])
doc.add_paragraph()

# ── Marco de priorización ─────────────────────────────────────────────────────
h(doc, '6a.1 Marco de Priorización de Mejoras', level=2)
p(doc,
  'Las mejoras se priorizan según una matriz de criticidad que combina el nivel '
  'de riesgo detectado por las herramientas con el impacto potencial sobre la '
  'seguridad, disponibilidad e integridad del sistema Eco-Alerta y los datos '
  'ciudadanos que gestiona.',
  space_after=6)

tp = doc.add_table(rows=1, cols=4)
tp.style = 'Table Grid'
tp.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(tp, ['Prioridad','Criterio','Tiempo Máximo de Resolución','Color'],
           [2.5,6.0,4.0,2.0])
for vals, bg in [
    (('P1 — Crítica',  'BLOCKER SonarQube / Riesgo Alto ZAP. Impacto directo en seguridad o integridad de datos.', 'Inmediato (≤ 7 días)',   'ROJO'),   'FFE0E0'),
    (('P2 — Alta',     'HIGH SonarQube / Riesgo Medio ZAP. Vulnerabilidades explotables bajo condiciones específicas.','Corto plazo (≤ 30 días)','NARANJA'),'FFF0DC'),
    (('P3 — Media',    'MAJOR/MEDIUM. Debilidades que afectan confiabilidad o generan deuda técnica significativa.','Mediano plazo (≤ 90 días)','AMARILLO'),'FFFFF0'),
    (('P4 — Baja',     'MINOR/LOW/INFO. Mejoras de calidad y mantenibilidad del código.','Largo plazo (≤ 180 días)','VERDE'),                              'F0FFF0'),
]:
    data_row(tp, vals, [2.5,6.0,4.0,2.0], row_bg=bg, sizes=[9,9,9,9], bold_cols=[0])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ÁREAS DE MEJORA
# ══════════════════════════════════════════════════════════════════════════════
h(doc, '6a.2 Plan de Mejoras por Área', level=2)
p(doc,
  'El plan se organiza en cinco áreas de mejora, agrupando los hallazgos por '
  'naturaleza técnica y alineándolos con los procesos definidos en los '
  'estándares ISO/IEC 12207 e ISO/IEC 14764.',
  space_after=8)

# ══════════════════════════════════════════════════════════════════════════════
# ÁREA 1: CREDENCIALES HARDCODEADAS
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Área 1: Gestión de Credenciales y Secretos en Código Fuente', level=3)
p(doc,
  'Hallazgos relacionados: SQ-SEC-01 al SQ-SEC-07 (7 vulnerabilidades, '
  '3 BLOCKER + 4 MAJOR). Estándar aplicado: ISO/IEC 12207 — Proceso de '
  'Implementación del Software (6.4.2) e ISO/IEC 25040 — Característica de '
  'Seguridad (Confidencialidad).',
  space_after=6)

p(doc,
  'Problema identificado: SonarQube detectó contraseñas y tokens de acceso '
  'escritos directamente en el código fuente Python (archivos create_admin_user.py, '
  'load_initial_data.py, make_inspector_admin.py y views.py). Cualquier persona '
  'con acceso al repositorio puede obtener credenciales de administrador y tokens '
  'de API, comprometiendo completamente la seguridad del sistema.',
  space_after=6)

ta1 = doc.add_table(rows=1, cols=5)
ta1.style = 'Table Grid'
ta1.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ta1, ['Acción','Descripción Técnica','Estándar','Prioridad','Responsable'],
           [3.0,6.5,2.5,1.5,2.5], bg='8B0000')

for vals in [
    ('A1.1 — Eliminar credenciales del código',
     'Remover todas las contraseñas y tokens hardcodeados de los archivos '
     'create_admin_user.py, load_initial_data.py, make_inspector_admin.py '
     'y views.py. Revocar y regenerar inmediatamente los tokens expuestos.',
     'ISO/IEC 12207 §6.4.2\nISO/IEC 25040 §7.3',
     'P1 — Crítica', 'Desarrollador Backend'),
    ('A1.2 — Implementar variables de entorno',
     'Reemplazar todas las credenciales por variables de entorno usando '
     'python-decouple o django-environ. Crear archivo .env con las variables '
     'y agregar .env al .gitignore para evitar su exposición en el repositorio.',
     'ISO/IEC 12207 §6.4.2\nISO/IEC 14764 §6.4',
     'P1 — Crítica', 'Desarrollador Backend'),
    ('A1.3 — Implementar escaneo automático de secretos',
     'Integrar git-secrets o detect-secrets como hook de pre-commit para '
     'bloquear automáticamente cualquier commit que contenga contraseñas '
     'o tokens en el código fuente.',
     'ISO/IEC 12207 §6.4.5\nISO/IEC 14764 §6.5',
     'P2 — Alta', 'DevOps / Desarrollador'),
    ('A1.4 — Auditoría del historial de Git',
     'Usar git filter-branch o BFG Repo Cleaner para eliminar del historial '
     'de commits los archivos que contenían credenciales, evitando que sean '
     'recuperables desde versiones anteriores del repositorio.',
     'ISO/IEC 14764 §6.4',
     'P2 — Alta', 'Desarrollador Backend'),
]:
    data_row(ta1, vals, [3.0,6.5,2.5,1.5,2.5], sizes=[9,9,8,9,9])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ÁREA 2: CONFIGURACIÓN DE SEGURIDAD HTTP
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Área 2: Configuración de Cabeceras y Políticas de Seguridad HTTP', level=3)
p(doc,
  'Hallazgos relacionados: ZAP-B-01, ZAP-F-01 (CSP), ZAP-B-02, ZAP-B-03, '
  'ZAP-F-05, ZAP-F-06 (Cookies), ZAP-F-04 (Clickjacking), ZAP-F-09 (MIME), '
  'ZAP-F-10 (HSTS). Total: 10 hallazgos. Estándar aplicado: ISO/IEC 25040 — '
  'Característica de Seguridad (Integridad, Confidencialidad, No repudio) e '
  'ISO/IEC 14764 — Mantenimiento Perfectivo.',
  space_after=6)

p(doc,
  'Problema identificado: El servidor Django no configura las cabeceras HTTP '
  'de seguridad recomendadas por OWASP. La cookie lab-session-id se emite sin '
  'los atributos HttpOnly ni SameSite, exponiendo la sesión a robo mediante '
  'JavaScript y ataques CSRF. La ausencia de Content-Security-Policy habilita '
  'ataques de inyección de scripts (XSS).',
  space_after=6)

ta2 = doc.add_table(rows=1, cols=5)
ta2.style = 'Table Grid'
ta2.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ta2, ['Acción','Descripción Técnica','Estándar','Prioridad','Responsable'],
           [3.0,6.5,2.5,1.5,2.5], bg='1A4B8C')

for vals in [
    ('A2.1 — Configurar Content-Security-Policy',
     'Agregar middleware CSP en Django settings.py usando django-csp. '
     'Definir política restrictiva: default-src \'self\'; script-src \'self\'; '
     'style-src \'self\' \'unsafe-inline\'; img-src \'self\' data:. '
     'Validar que no rompa funcionalidad del frontend React.',
     'ISO/IEC 25040 §7.3\nOWASP A05:2021',
     'P2 — Alta', 'Desarrollador Backend'),
    ('A2.2 — Asegurar atributos de cookies',
     'En Django settings.py configurar: SESSION_COOKIE_HTTPONLY = True, '
     'SESSION_COOKIE_SAMESITE = \'Strict\', SESSION_COOKIE_SECURE = True '
     '(para producción con HTTPS). Aplicar mismos atributos a la cookie '
     'lab-session-id identificada por ZAP.',
     'ISO/IEC 14764 §6.4\nCWE-1004, CWE-1275',
     'P2 — Alta', 'Desarrollador Backend'),
    ('A2.3 — Configurar cabeceras anti-clickjacking',
     'Activar X_FRAME_OPTIONS = \'DENY\' en Django settings.py. '
     'Alternativamente incluir frame-ancestors \'none\' en la política CSP. '
     'Verificar que el middleware SecurityMiddleware está activo.',
     'ISO/IEC 25040 §7.3\nOWASP A05:2021',
     'P2 — Alta', 'Desarrollador Backend'),
    ('A2.4 — Configurar HSTS y X-Content-Type-Options',
     'Activar SECURE_HSTS_SECONDS = 31536000 y SECURE_HSTS_INCLUDE_SUBDOMAINS = True '
     'para producción. Agregar header X-Content-Type-Options: nosniff en todas '
     'las respuestas para prevenir MIME-sniffing.',
     'ISO/IEC 14764 §6.4\nOWASP A05:2021',
     'P2 — Alta', 'Desarrollador Backend'),
    ('A2.5 — Desactivar modo DEBUG en producción',
     'Configurar DEBUG = False en el entorno de producción usando variables '
     'de entorno. El modo DEBUG expone rutas internas y configuración del '
     'proyecto en mensajes de error, detectado en el escaneo ZAP.',
     'ISO/IEC 12207 §6.4.2\nISO/IEC 14764 §6.3',
     'P1 — Crítica', 'Desarrollador Backend'),
]:
    data_row(ta2, vals, [3.0,6.5,2.5,1.5,2.5], sizes=[9,9,8,9,9])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ÁREA 3: CONTROL DE ACCESO A ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Área 3: Control de Acceso y Restricción de Métodos HTTP', level=3)
p(doc,
  'Hallazgos relacionados: SQ-SEC-08, SQ-SEC-09 (métodos HTTP no restringidos), '
  'ZAP-F-02 (CORS), ZAP-F-03 (Directory Browsing), ZAP-F-12 (token en URL). '
  'Estándar aplicado: ISO/IEC 25040 — Característica de Seguridad (Control de '
  'Acceso) e ISO/IEC 12207 — Proceso de Verificación.',
  space_after=6)

ta3 = doc.add_table(rows=1, cols=5)
ta3.style = 'Table Grid'
ta3.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ta3, ['Acción','Descripción Técnica','Estándar','Prioridad','Responsable'],
           [3.0,6.5,2.5,1.5,2.5], bg='1A4B8C')

for vals in [
    ('A3.1 — Restringir métodos HTTP en endpoints',
     'En urls.py de Django, reemplazar path() genérico por vistas basadas '
     'en clases con http_method_names definido, o usar el decorador '
     '@require_http_methods([\'GET\', \'POST\']) según corresponda a cada '
     'endpoint de la API REST.',
     'ISO/IEC 12207 §6.4.5\nSonarQube python:S3752',
     'P2 — Alta', 'Desarrollador Backend'),
    ('A3.2 — Restringir configuración CORS',
     'En Django settings.py (django-cors-headers): reemplazar '
     'CORS_ALLOW_ALL_ORIGINS = True por CORS_ALLOWED_ORIGINS con lista '
     'explícita de dominios autorizados. En producción solo el dominio del '
     'frontend oficial.',
     'ISO/IEC 25040 §7.3\nOWASP A01:2021',
     'P2 — Alta', 'Desarrollador Backend'),
    ('A3.3 — Deshabilitar directory browsing',
     'Configurar el servidor web (nginx/Apache en producción) para denegar '
     'el listado de directorios. En el servidor Vite de desarrollo, asegurarse '
     'que node_modules no sea accesible públicamente.',
     'ISO/IEC 14764 §6.4\nOWASP A05:2021',
     'P2 — Alta', 'DevOps / Desarrollador'),
    ('A3.4 — Eliminar tokens en URL',
     'Modificar el frontend React para que los tokens JWT se transmitan '
     'exclusivamente mediante el header Authorization: Bearer <token>. '
     'Nunca incluir tokens como parámetros querystring (?token=...).',
     'ISO/IEC 12207 §6.4.2\nCWE-200',
     'P2 — Alta', 'Desarrollador Frontend'),
]:
    data_row(ta3, vals, [3.0,6.5,2.5,1.5,2.5], sizes=[9,9,8,9,9])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ÁREA 4: CONFIABILIDAD DEL CÓDIGO
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Área 4: Mejoras de Confiabilidad del Código Fuente', level=3)
p(doc,
  'Hallazgos relacionados: 229 issues de Reliability detectados por SonarQube '
  '(44 High + 185 Medium/Low). Estándar aplicado: ISO/IEC 25040 — '
  'Característica de Confiabilidad (Madurez, Tolerancia a Fallos, '
  'Recuperabilidad) e ISO/IEC 14764 — Mantenimiento Correctivo y Preventivo.',
  space_after=6)

ta4 = doc.add_table(rows=1, cols=5)
ta4.style = 'Table Grid'
ta4.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ta4, ['Acción','Descripción Técnica','Estándar','Prioridad','Responsable'],
           [3.0,6.5,2.5,1.5,2.5], bg='1A4B8C')

for vals in [
    ('A4.1 — Refactorizar manejo de excepciones',
     'Revisar todos los bloques try/except en el backend Python (views.py, '
     'services.py). Reemplazar except genéricos por excepciones específicas. '
     'Implementar logging adecuado con el módulo logging de Django para '
     'registrar errores sin exponer información sensible.',
     'ISO/IEC 14764 §6.4\nISO/IEC 25040 §7.2',
     'P2 — Alta', 'Desarrollador Backend'),
    ('A4.2 — Resolver issues High de Reliability',
     'Priorizar los 44 issues de severidad High identificados por SonarQube. '
     'Enfocar en variables posiblemente no inicializadas y uso incorrecto de '
     'APIs. Usar el dashboard de SonarQube como guía de trabajo sprint a sprint.',
     'ISO/IEC 14764 §6.4\nISO/IEC 12207 §6.4.5',
     'P2 — Alta', 'Equipo de Desarrollo'),
    ('A4.3 — Corregir uso de APIs deprecadas en React',
     'Identificar y actualizar los componentes React que utilizan APIs '
     'deprecadas. Actualizar dependencias del frontend a versiones estables '
     'actuales. Ejecutar npm audit para detectar vulnerabilidades conocidas.',
     'ISO/IEC 14764 §6.4.2\nISO/IEC 12207 §6.4.5',
     'P3 — Media', 'Desarrollador Frontend'),
    ('A4.4 — Implementar pruebas unitarias',
     'Crear suite de pruebas unitarias para los módulos críticos del backend '
     '(views.py, permissions.py, notification_service.py) usando pytest-django. '
     'Meta: cobertura mínima del 70% en módulos de seguridad.',
     'ISO/IEC 12207 §6.4.6\nISO/IEC 25040 §7.5',
     'P3 — Media', 'Equipo de Desarrollo'),
]:
    data_row(ta4, vals, [3.0,6.5,2.5,1.5,2.5], sizes=[9,9,8,9,9])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ÁREA 5: MANTENIBILIDAD Y DEUDA TÉCNICA
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Área 5: Reducción de Deuda Técnica y Mejoras de Mantenibilidad', level=3)
p(doc,
  'Hallazgos relacionados: 458 issues de Maintainability detectados por '
  'SonarQube. Estándar aplicado: ISO/IEC 25040 — Característica de '
  'Mantenibilidad (Analizabilidad, Modificabilidad, Reusabilidad) e '
  'ISO/IEC 14764 — Mantenimiento Perfectivo.',
  space_after=6)

ta5 = doc.add_table(rows=1, cols=5)
ta5.style = 'Table Grid'
ta5.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ta5, ['Acción','Descripción Técnica','Estándar','Prioridad','Responsable'],
           [3.0,6.5,2.5,1.5,2.5], bg='1A4B8C')

for vals in [
    ('A5.1 — Eliminar código duplicado',
     'Extraer el literal redis://127.0.0.1:6379/1 (repetido 3 veces) a una '
     'constante en settings.py o variable de entorno REDIS_URL. Aplicar el '
     'principio DRY (Don\'t Repeat Yourself) en todo el backend.',
     'ISO/IEC 25040 §7.4\nISO/IEC 14764 §6.4.3',
     'P3 — Media', 'Desarrollador Backend'),
    ('A5.2 — Documentar funciones críticas',
     'Agregar docstrings a todas las funciones públicas en views.py, '
     'permissions.py y services.py del backend. Para el frontend, '
     'documentar los componentes React principales usando JSDoc.',
     'ISO/IEC 14764 §6.4.3\nISO/IEC 12207 §6.4.3',
     'P3 — Media', 'Equipo de Desarrollo'),
    ('A5.3 — Eliminar imports y variables no usados',
     'Ejecutar flake8 y pylint en el backend para identificar y eliminar '
     'imports no utilizados. En el frontend usar ESLint con regla '
     'no-unused-vars activa. Meta: cero advertencias de lint en CI/CD.',
     'ISO/IEC 25040 §7.4\nISO/IEC 14764 §6.4.3',
     'P3 — Media', 'Equipo de Desarrollo'),
    ('A5.4 — Refactorizar funciones de alta complejidad',
     'Identificar las funciones de views.py con mayor complejidad cognitiva '
     'según SonarQube. Dividir funciones que superen 15 líneas de lógica '
     'en sub-funciones con responsabilidad única (principio SRP).',
     'ISO/IEC 14764 §6.4.3\nISO/IEC 25040 §7.4',
     'P4 — Baja', 'Desarrollador Backend'),
    ('A5.5 — Integrar análisis continuo en pipeline',
     'Configurar SonarQube para ejecutarse automáticamente en cada pull '
     'request mediante GitHub Actions o similar. Establecer Quality Gate '
     'que rechace código con nuevos issues de seguridad o Blocker.',
     'ISO/IEC 12207 §6.4.5\nISO/IEC 14764 §6.5',
     'P3 — Media', 'DevOps / Equipo'),
]:
    data_row(ta5, vals, [3.0,6.5,2.5,1.5,2.5], sizes=[9,9,8,9,9])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6a.3 ROADMAP DE IMPLEMENTACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h(doc, '6a.3 Roadmap de Implementación', level=2)
p(doc,
  'La siguiente tabla presenta el cronograma de implementación de las mejoras, '
  'ordenado por prioridad y agrupado en sprints de desarrollo. Se alinea con '
  'ISO/IEC 12207 §6.4 (Proceso de Implementación) e ISO/IEC 14764 §6.4 '
  '(Implementación del Mantenimiento).',
  space_after=6)

tr = doc.add_table(rows=1, cols=5)
tr.style = 'Table Grid'
tr.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(tr, ['Sprint / Fase','Período','Acciones','Criterio de Éxito','Estándar'],
           [2.5,2.5,4.0,5.0,2.0])

roadmap_data = [
    ('Sprint 0\n(Urgente)',
     'Semana 1–2',
     'A1.1, A1.2, A2.5\n(Credenciales + DEBUG)',
     'Cero credenciales en código. Variables de entorno activas. DEBUG=False en producción. SonarQube sin BLOCKER.',
     'ISO/IEC 14764\n§6.3 y §6.4',
     'FFE0E0'),
    ('Sprint 1\n(Seguridad HTTP)',
     'Semana 3–6',
     'A2.1, A2.2, A2.3, A2.4\nA3.1, A3.2, A3.3, A3.4',
     'ZAP re-scan sin alertas de nivel Medio. CSP activa. Cookies con HttpOnly+SameSite. CORS restringido.',
     'ISO/IEC 25040\n§7.3',
     'FFF0DC'),
    ('Sprint 2\n(Confiabilidad)',
     'Mes 2',
     'A1.3, A1.4, A4.1, A4.2',
     'Issues High de SonarQube reducidos en ≥80%. Cobertura de pruebas ≥50% en módulos críticos.',
     'ISO/IEC 12207\n§6.4.5–6',
     'FFFFF0'),
    ('Sprint 3\n(Calidad)',
     'Mes 3',
     'A4.3, A4.4, A5.1, A5.2, A5.3',
     'SonarQube Quality Gate en verde. Cero issues de Reliability High pendientes. Documentación actualizada.',
     'ISO/IEC 14764\n§6.4.3',
     'F0FFF0'),
    ('Sprint 4\n(Mejora continua)',
     'Mes 4–6',
     'A5.4, A5.5\n(Deuda técnica + CI/CD)',
     'Pipeline CI/CD con SonarQube integrado. Quality Gate bloqueante activo. Deuda técnica reducida en ≥50%.',
     'ISO/IEC 12207\n§6.4.5',
     'F0F0FF'),
]
for vals in roadmap_data:
    row = tr.add_row()
    for i, (val, w) in enumerate(zip(vals[:5], [2.5,2.5,4.0,5.0,2.0])):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, vals[5])
        set_borders(cell)
        par = cell.paragraphs[0]
        run = par.add_run(val)
        run.font.size = Pt(8.5)
        run.font.name = 'Arial'
        if i == 0:
            run.font.bold = True
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6a.4 ALINEACIÓN CON ESTÁNDARES
# ══════════════════════════════════════════════════════════════════════════════
h(doc, '6a.4 Alineación Explícita con Estándares ISO', level=2)
p(doc,
  'La siguiente tabla muestra cómo cada área de mejora se vincula '
  'específicamente con los procesos y características definidas en los '
  'estándares internacionales aplicados.',
  space_after=6)

te = doc.add_table(rows=1, cols=4)
te.style = 'Table Grid'
te.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(te, ['Área de Mejora','ISO/IEC 25040','ISO/IEC 12207','ISO/IEC 14764'],
           [4.0,4.0,4.0,4.0])

for vals in [
    ('Área 1: Gestión de Credenciales',
     '§7.3 Seguridad — Confidencialidad: proteger datos contra acceso no autorizado',
     '§6.4.2 Implementación: el software no debe contener información sensible en texto plano',
     '§6.4 Impl. del Mantenimiento: corrección de defectos de seguridad como mantenimiento correctivo'),
    ('Área 2: Cabeceras HTTP',
     '§7.3 Seguridad — Integridad: garantizar que los datos no sean alterados en tránsito',
     '§6.4.5 Integración del Software: verificar que componentes del sistema cumplen requisitos de seguridad',
     '§6.4.3 Mantenimiento Perfectivo: mejoras que aumentan el rendimiento y seguridad sin corregir defectos'),
    ('Área 3: Control de Acceso',
     '§7.3 Seguridad — Control de Acceso: solo usuarios autorizados pueden ejecutar operaciones',
     '§6.4.6 Pruebas de Calificación: verificar que los controles de acceso funcionan correctamente',
     '§6.4 Mantenimiento Correctivo: corrección de vulnerabilidades de control de acceso identificadas'),
    ('Área 4: Confiabilidad',
     '§7.2 Confiabilidad — Tolerancia a Fallos: el sistema mantiene operación ante condiciones adversas',
     '§6.4.6 Pruebas: validar que el software maneja correctamente condiciones de error',
     '§6.4.2 Análisis del Problema: identificar la causa raíz de defectos de confiabilidad'),
    ('Área 5: Mantenibilidad',
     '§7.4 Mantenibilidad — Modificabilidad: facilidad para realizar cambios sin introducir defectos',
     '§6.4.3 Diseño de la Arquitectura: código bien estructurado y documentado según estándares',
     '§6.4.3 Impl. de la Modificación: aplicar cambios de mantenimiento perfectivo de forma controlada'),
]:
    data_row(te, vals, [4.0,4.0,4.0,4.0], sizes=[9,9,9,9])
doc.add_paragraph()

# ── Cierre de sección ─────────────────────────────────────────────────────────
p(doc,
  'El plan de mejoras diseñado aborda la totalidad de los hallazgos detectados '
  'por las herramientas automatizadas, priorizando la eliminación inmediata de '
  'las vulnerabilidades críticas (credenciales hardcodeadas y modo DEBUG activo) '
  'antes de avanzar hacia mejoras de confiabilidad y mantenibilidad. La '
  'alineación con ISO/IEC 25040, ISO/IEC 12207 e ISO/IEC 14764 garantiza que '
  'las mejoras se implementen siguiendo procesos formales y verificables, '
  'contribuyendo al ciclo de vida de calidad del software Eco-Alerta.',
  italic=True, size=10, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
out = '/sessions/tender-gifted-cray/mnt/ecoalerta/seccion6a_v1.docx'
doc.save(out)
print(f'OK: {out}')
