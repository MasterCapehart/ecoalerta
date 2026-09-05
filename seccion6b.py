#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera seccion6b_v1.docx — Sección 6b: Resumen del Plan de Acción (Matriz de Control)"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.0)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)

for h_name, sz in [('Heading 1',14),('Heading 2',12),('Heading 3',11)]:
    s = doc.styles[h_name]
    s.font.name  = 'Arial'
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = RGBColor(0,0,0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h(doc, text, level=1):
    par = doc.add_heading(text, level=level)
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return par

def p(doc, text, bold=False, italic=False, size=11, space_after=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after  = Pt(space_after)
    par.paragraph_format.space_before = Pt(2)
    run = par.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    return par

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'AAAAAA')
        tcB.append(b)
    tcPr.append(tcB)

def header_row(table, headers, widths, bg='1F5C99'):
    row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, bg)
        set_borders(cell)
        par = cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(hdr)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size      = Pt(9)
        run.font.name      = 'Arial'

def add_row(table, values, widths, bg=None, sizes=None, bold_cols=None, center_cols=None):
    row = table.add_row()
    for i, (val, w) in enumerate(zip(values, widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        if bg:
            set_cell_bg(cell, bg)
        set_borders(cell)
        par = cell.paragraphs[0]
        if center_cols and i in center_cols:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sz  = sizes[i] if sizes else 9
        run = par.add_run(str(val))
        run.font.size = Pt(sz)
        run.font.name = 'Arial'
        if bold_cols and i in bold_cols:
            run.font.bold = True
    return row

def priority_bg(p):
    return {
        'P1':'FFCCCC', 'P2':'FFE0B2',
        'P3':'FFF9C4', 'P4':'DCEDC8',
    }.get(p[:2], 'FFFFFF')

def estado_bg(e):
    return {
        'Pendiente':'FFE0E0',
        'En Progreso':'FFF3CD',
        'Completado':'D4EDDA',
    }.get(e, 'FFFFFF')

# ══════════════════════════════════════════════════════════════════════════════
# TÍTULO
# ══════════════════════════════════════════════════════════════════════════════
h(doc, '6b. Resumen del Plan de Acción — Matriz de Control', level=2)

p(doc,
  'La Matriz de Control resume todas las acciones de mejora definidas en el '
  'plan, estableciendo para cada una su identificador, hallazgo de origen, '
  'prioridad, responsable, plazo de ejecución, criterio de verificación y '
  'estado inicial. Esta matriz permite el seguimiento y control del avance '
  'del plan de mejoras a lo largo del tiempo.',
  space_after=10)

# ══════════════════════════════════════════════════════════════════════════════
# LEYENDA DE PRIORIDADES
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Leyenda de Prioridades y Estados', level=3)

tl = doc.add_table(rows=1, cols=6)
tl.style = 'Table Grid'
tl.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(tl,
    ['Prioridad','Descripción','Color','Estado','Descripción','Color'],
    [1.5, 4.5, 1.5, 2.0, 3.5, 1.5])

for vals, p_bg, e_bg in [
    (('P1 — Crítica', 'BLOCKER / impacto directo en seguridad', '', 'Pendiente',    'No iniciado aún',          ''),
     'FFCCCC', 'FFE0E0'),
    (('P2 — Alta',    'Riesgo Medio / explotable en condiciones específicas', '', 'En Progreso', 'Actualmente en desarrollo', ''),
     'FFE0B2', 'FFF3CD'),
    (('P3 — Media',   'Confiabilidad o deuda técnica significativa', '', 'Completado',  'Implementado y verificado', ''),
     'FFF9C4', 'D4EDDA'),
    (('P4 — Baja',    'Mejoras de calidad y mantenibilidad', '', '', '', ''),
     'DCEDC8', 'FFFFFF'),
]:
    row = tl.add_row()
    for i, w in enumerate([1.5,4.5,1.5,2.0,3.5,1.5]):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_borders(cell)
        if i == 2:
            set_cell_bg(cell, p_bg)
        elif i == 5 and e_bg != 'FFFFFF':
            set_cell_bg(cell, e_bg)
        par = cell.paragraphs[0]
        run = par.add_run(vals[i])
        run.font.size = Pt(8.5)
        run.font.name = 'Arial'
        if i in [0, 3]:
            run.font.bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# MATRIZ DE CONTROL PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Matriz de Control de Mejoras', level=3)

p(doc,
  'La siguiente matriz consolida las 22 acciones de mejora definidas en el '
  'plan, organizadas por área y prioridad.',
  space_after=6)

widths = [1.0, 1.5, 3.5, 4.5, 2.0, 2.0, 3.5, 2.0]
tm = doc.add_table(rows=1, cols=8)
tm.style = 'Table Grid'
tm.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(tm,
    ['ID Acción', 'ID Hallazgo', 'Acción de Mejora',
     'Criterio de Verificación', 'Responsable', 'Plazo', 'Estándar Aplicado', 'Estado'],
    widths, bg='1F5C99')

acciones = [
    # ── ÁREA 1: Credenciales ──────────────────────────────────────────────────
    ('A1.1', 'SQ-SEC-01\nSQ-SEC-02\nSQ-SEC-03',
     'Eliminar credenciales y tokens hardcodeados del código fuente y revocar los expuestos',
     'SonarQube sin issues BLOCKER de tipo S6437 y S6418. Tokens revocados y regenerados.',
     'Desarrollador Backend', 'Semana 1',
     'ISO/IEC 12207\nProceso de Implementación\nISO/IEC 14764\nMantenimiento Correctivo',
     'Pendiente', 'P1'),
    ('A1.2', 'SQ-SEC-01\nSQ-SEC-04\nSQ-SEC-07',
     'Implementar variables de entorno con python-decouple. Agregar .env al .gitignore',
     'Archivo .env creado. Ninguna credencial en texto plano en el código. .env en .gitignore.',
     'Desarrollador Backend', 'Semana 1',
     'ISO/IEC 12207\nProceso de Implementación\nISO/IEC 14764\nMantenimiento Correctivo',
     'Pendiente', 'P1'),
    ('A1.3', 'SQ-SEC-01\nal SQ-SEC-07',
     'Integrar git-secrets como hook de pre-commit para bloquear commits con credenciales',
     'Pre-commit hook activo. Prueba de commit con credencial ficticia es rechazado automáticamente.',
     'DevOps', 'Semana 3',
     'ISO/IEC 12207\nProceso de Verificación\nISO/IEC 14764\nMantenimiento Preventivo',
     'Pendiente', 'P2'),
    ('A1.4', 'SQ-SEC-01\nal SQ-SEC-07',
     'Limpiar historial de Git con BFG Repo Cleaner para eliminar commits con credenciales',
     'git log no contiene archivos con credenciales en ningún commit anterior.',
     'Desarrollador Backend', 'Semana 4',
     'ISO/IEC 14764\nMantenimiento Correctivo',
     'Pendiente', 'P2'),

    # ── ÁREA 2: Cabeceras HTTP ────────────────────────────────────────────────
    ('A2.1', 'ZAP-B-01\nZAP-F-01',
     'Configurar Content-Security-Policy en Django usando django-csp',
     'Re-escaneo ZAP sin alerta de CSP. Header Content-Security-Policy presente en todas las respuestas.',
     'Desarrollador Backend', 'Semana 3',
     'ISO/IEC 25040\nCaract. Seguridad\nISO/IEC 14764\nMantenimiento Perfectivo',
     'Pendiente', 'P2'),
    ('A2.2', 'ZAP-B-02\nZAP-B-03\nZAP-F-05\nZAP-F-06',
     'Configurar SESSION_COOKIE_HTTPONLY=True y SESSION_COOKIE_SAMESITE="Strict" en Django',
     'Re-escaneo ZAP sin alertas CWE-1004 ni CWE-1275. Cookie lab-session-id con HttpOnly y SameSite.',
     'Desarrollador Backend', 'Semana 3',
     'ISO/IEC 25040\nCaract. Seguridad\nISO/IEC 14764\nMantenimiento Perfectivo',
     'Pendiente', 'P2'),
    ('A2.3', 'ZAP-F-04',
     'Configurar X_FRAME_OPTIONS="DENY" en Django settings.py',
     'Re-escaneo ZAP sin alerta de Clickjacking. Header X-Frame-Options presente en respuestas.',
     'Desarrollador Backend', 'Semana 3',
     'ISO/IEC 25040\nCaract. Seguridad',
     'Pendiente', 'P2'),
    ('A2.4', 'ZAP-F-09\nZAP-F-10',
     'Activar HSTS y X-Content-Type-Options: nosniff en todas las respuestas HTTP',
     'Headers HSTS y X-Content-Type-Options presentes. ZAP sin alertas relacionadas.',
     'Desarrollador Backend', 'Semana 4',
     'ISO/IEC 25040\nCaract. Seguridad\nISO/IEC 14764\nMantenimiento Perfectivo',
     'Pendiente', 'P2'),
    ('A2.5', 'ZAP-B-01\nZAP-F-01',
     'Desactivar DEBUG=True en entorno de producción usando variables de entorno',
     'Django no expone rutas internas en mensajes de error 404. DEBUG=False verificado.',
     'Desarrollador Backend', 'Semana 1',
     'ISO/IEC 12207\nProceso de Implementación',
     'Pendiente', 'P1'),

    # ── ÁREA 3: Control de Acceso ─────────────────────────────────────────────
    ('A3.1', 'SQ-SEC-08\nSQ-SEC-09',
     'Agregar decorador @require_http_methods en endpoints de urls.py sin restricción de método',
     'SonarQube sin issues python:S3752. Endpoints solo aceptan verbos HTTP declarados.',
     'Desarrollador Backend', 'Semana 4',
     'ISO/IEC 12207\nProceso de Verificacion\nISO/IEC 14764\nMantenimiento Correctivo',
     'Pendiente', 'P2'),
    ('A3.2', 'ZAP-F-02',
     'Restringir CORS_ALLOWED_ORIGINS a dominios autorizados en Django settings.py',
     'CORS_ALLOW_ALL_ORIGINS = False. Solo el dominio del frontend en lista blanca.',
     'Desarrollador Backend', 'Semana 4',
     'ISO/IEC 25040\nCaract. Seguridad\nISO/IEC 12207\nProceso de Verificacion',
     'Pendiente', 'P2'),
    ('A3.3', 'ZAP-F-03',
     'Deshabilitar directory browsing en servidor web. Bloquear acceso a /node_modules/',
     'GET /node_modules/ retorna 403 Forbidden. ZAP sin alerta de Directory Browsing.',
     'DevOps', 'Semana 4',
     'ISO/IEC 14764\nMantenimiento Preventivo',
     'Pendiente', 'P2'),
    ('A3.4', 'ZAP-F-12',
     'Modificar frontend React para transmitir JWT solo en header Authorization, nunca en URL',
     'Ninguna URL del frontend contiene el parámetro ?token=. JWT solo en Authorization header.',
     'Desarrollador Frontend', 'Semana 4',
     'ISO/IEC 12207\nProceso de Implementacion\nISO/IEC 25040\nCaract. Seguridad',
     'Pendiente', 'P2'),

    # ── ÁREA 4: Confiabilidad ─────────────────────────────────────────────────
    ('A4.1', 'SonarQube\nReliability High',
     'Refactorizar bloques try/except genéricos en backend Python. Implementar logging apropiado',
     'SonarQube Reliability High reducido en 80%. Logs de error configurados en Django.',
     'Desarrollador Backend', 'Mes 2',
     'ISO/IEC 14764\nMantenimiento Correctivo\nISO/IEC 25040\nCaract. Confiabilidad',
     'Pendiente', 'P2'),
    ('A4.2', 'SonarQube\n44 issues High',
     'Resolver los 44 issues de severidad High de Reliability priorizados en SonarQube',
     'Dashboard SonarQube con 0 issues High de Reliability pendientes.',
     'Equipo de Desarrollo', 'Mes 2',
     'ISO/IEC 14764\nMantenimiento Correctivo\nISO/IEC 12207\nProceso de Verificacion',
     'Pendiente', 'P2'),
    ('A4.3', 'SonarQube\nFrontend Issues',
     'Actualizar APIs deprecadas de React y ejecutar npm audit para resolver vulnerabilidades',
     'npm audit sin vulnerabilidades críticas. Sin warnings de APIs deprecadas en consola.',
     'Desarrollador Frontend', 'Mes 2',
     'ISO/IEC 14764\nMantenimiento Adaptativo\nISO/IEC 12207\nProceso de Implementacion',
     'Pendiente', 'P3'),
    ('A4.4', 'SonarQube\nConfiabilidad',
     'Crear suite de pruebas unitarias con pytest-django. Meta: 70% cobertura en módulos críticos',
     'Reporte pytest con cobertura mayor o igual al 70% en views.py, permissions.py y services.py.',
     'Equipo de Desarrollo', 'Mes 2-3',
     'ISO/IEC 12207\nProceso de Pruebas\nISO/IEC 25040\nEvaluacion de Calidad',
     'Pendiente', 'P3'),

    # ── ÁREA 5: Mantenibilidad ────────────────────────────────────────────────
    ('A5.1', 'SonarQube\nMaintainability',
     'Extraer literal redis://127.0.0.1:6379/1 a constante en settings.py o variable de entorno',
     'SonarQube sin issue de duplicación del literal Redis. Constante REDIS_URL definida y usada.',
     'Desarrollador Backend', 'Mes 3',
     'ISO/IEC 25040\nCaract. Mantenibilidad\nISO/IEC 14764\nMantenimiento Perfectivo',
     'Pendiente', 'P3'),
    ('A5.2', 'SonarQube\n458 issues Maint.',
     'Agregar docstrings a funciones públicas de views.py, permissions.py y servicios del backend',
     'Cobertura de documentación mayor al 80% en módulos críticos del backend.',
     'Equipo de Desarrollo', 'Mes 3',
     'ISO/IEC 14764\nMantenimiento Perfectivo\nISO/IEC 12207\nDocumentacion',
     'Pendiente', 'P3'),
    ('A5.3', 'SonarQube\nMaintainability',
     'Eliminar imports y variables no usados con flake8 (backend) y ESLint (frontend)',
     'Cero warnings de flake8 y ESLint en modo CI. Pipeline no acepta código con linting fallido.',
     'Equipo de Desarrollo', 'Mes 3',
     'ISO/IEC 25040\nCaract. Mantenibilidad\nISO/IEC 14764\nMantenimiento Perfectivo',
     'Pendiente', 'P3'),
    ('A5.4', 'SonarQube\nComplejidad',
     'Refactorizar funciones de alta complejidad cognitiva en views.py aplicando principio SRP',
     'SonarQube sin issues de complejidad cognitiva mayor a 15. Funciones con responsabilidad única.',
     'Desarrollador Backend', 'Mes 4-5',
     'ISO/IEC 25040\nCaract. Mantenibilidad\nISO/IEC 14764\nMantenimiento Perfectivo',
     'Pendiente', 'P4'),
    ('A5.5', 'Todos los\nhallazgos',
     'Integrar SonarQube en pipeline CI/CD con Quality Gate bloqueante para nuevos issues',
     'Pipeline activo en GitHub Actions. Quality Gate rechaza PRs con nuevos BLOCKER o High.',
     'DevOps', 'Mes 4',
     'ISO/IEC 12207\nProceso de Verificacion\nISO/IEC 14764\nMantenimiento Preventivo',
     'Pendiente', 'P3'),
]

for accion in acciones:
    id_a, id_h, desc, criterio, resp, plazo, estandar, estado, prio = accion
    bg = priority_bg(prio)
    row = tm.add_row()
    vals = [id_a, id_h, desc, criterio, resp, plazo, estandar, estado]
    for i, (val, w) in enumerate(zip(vals, widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_borders(cell)
        # color fondo según prioridad en columna de acción
        if i == 0:
            set_cell_bg(cell, bg)
        elif i == 7:
            set_cell_bg(cell, estado_bg(estado))
        par = cell.paragraphs[0]
        if i in [4, 5, 7]:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(val)
        run.font.size = Pt(8 if i not in [0,1] else 8.5)
        run.font.name = 'Arial'
        if i == 0:
            run.font.bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# RESUMEN ESTADÍSTICO
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Resumen Estadístico del Plan', level=3)

ts = doc.add_table(rows=1, cols=5)
ts.style = 'Table Grid'
ts.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ts,
    ['Prioridad', 'N° Acciones', '% del Plan', 'Plazo Máximo', 'Área Principal'],
    [3.0, 2.0, 2.0, 3.0, 5.5])

for vals, bg in [
    (('P1 — Crítica',   '3',  '13,6%', 'Semana 1 (7 días)',   'Credenciales + DEBUG'),           'FFCCCC'),
    (('P2 — Alta',      '12', '54,5%', 'Semana 3-4 (30 días)','Cabeceras HTTP + Acceso + Confiabilidad'), 'FFE0B2'),
    (('P3 — Media',     '6',  '27,3%', 'Mes 2-4 (90 días)',   'Confiabilidad + Mantenibilidad'), 'FFF9C4'),
    (('P4 — Baja',      '1',  '4,5%',  'Mes 4-5 (180 días)',  'Refactorización avanzada'),        'DCEDC8'),
    (('TOTAL',          '22', '100%',  '—',                   '5 áreas de mejora'),               'E8E8E8'),
]:
    row = ts.add_row()
    for i, (val, w) in enumerate(zip(vals, [3.0,2.0,2.0,3.0,5.5])):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, bg)
        set_borders(cell)
        par = cell.paragraphs[0]
        if i in [1,2]:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(val)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        if i == 0:
            run.font.bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# INDICADORES DE SEGUIMIENTO
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Indicadores de Seguimiento del Plan', level=3)

p(doc,
  'Para verificar el avance y eficacia del plan de mejoras, se definen los '
  'siguientes indicadores clave de rendimiento (KPI) que deben ser medidos '
  'al finalizar cada sprint:',
  space_after=6)

ti = doc.add_table(rows=1, cols=4)
ti.style = 'Table Grid'
ti.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(ti,
    ['Indicador', 'Valor Actual\n(línea base)', 'Meta', 'Herramienta de Medición'],
    [5.0, 3.0, 4.0, 4.0])

indicadores = [
    ('Issues BLOCKER en SonarQube',
     '3 BLOCKER activos',
     '0 BLOCKER al finalizar Sprint 0',
     'SonarQube Dashboard'),
    ('Issues de Seguridad en SonarQube',
     '9 vulnerabilidades',
     '0 vulnerabilidades al finalizar Sprint 1',
     'SonarQube Dashboard'),
    ('Alertas de nivel Medio en OWASP ZAP',
     '5 alertas Medio (4 frontend + 1 backend)',
     '0 alertas Medio al finalizar Sprint 1',
     'Re-escaneo OWASP ZAP'),
    ('Alertas de nivel Bajo en OWASP ZAP',
     '10 alertas Bajo',
     'Reduccion del 80% al finalizar Sprint 1',
     'Re-escaneo OWASP ZAP'),
    ('Issues High de Reliability en SonarQube',
     '44 issues High',
     'Reduccion del 80% al finalizar Sprint 2',
     'SonarQube Dashboard'),
    ('Cobertura de pruebas unitarias (backend)',
     '0% (sin pruebas automatizadas)',
     'Mayor o igual a 70% en modulos criticos al Sprint 3',
     'pytest-cov / SonarQube'),
    ('Total de issues en SonarQube',
     '553 issues',
     'Reduccion del 60% (menos de 220) al Sprint 4',
     'SonarQube Dashboard'),
    ('Deuda tecnica estimada',
     'Alta (segun SonarQube)',
     'Quality Gate en estado Passed sin advertencias al Sprint 4',
     'SonarQube Quality Gate'),
]

for vals in indicadores:
    add_row(ti, vals, [5.0,3.0,4.0,4.0], sizes=[9,9,9,9])

doc.add_paragraph()

# ── Cierre ────────────────────────────────────────────────────────────────────
p(doc,
  'Esta Matriz de Control constituye el instrumento de seguimiento del Plan '
  'de Mejoras definido para el sistema Eco-Alerta. Su aplicacion sistematica, '
  'combinada con la re-ejecucion periodica de OWASP ZAP y SonarQube como '
  'herramientas de verificacion, garantiza que las mejoras implementadas '
  'sean efectivas y medibles, elevando progresivamente el nivel de seguridad, '
  'confiabilidad y mantenibilidad de la plataforma.',
  italic=True, size=10, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
out = '/sessions/tender-gifted-cray/mnt/ecoalerta/seccion6b_v1.docx'
doc.save(out)
print(f'OK: {out}')
