#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera seccion6b_v2.docx — Sección 6b: Resumen del Plan de Acción (Matriz de Control) formato ejecutivo"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.0)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)

for h_name, sz in [('Heading 1',14),('Heading 2',12),('Heading 3',11)]:
    s = doc.styles[h_name]
    s.font.name  = 'Arial'
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = RGBColor(0,0,0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h(doc, text, level=1):
    par = doc.add_heading(text, level=level)
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return par

def p(doc, text, bold=False, italic=False, size=11, space_after=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after  = Pt(space_after)
    par.paragraph_format.space_before = Pt(2)
    run = par.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    return par

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_borders(cell, color='AAAAAA'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcB.append(b)
    tcPr.append(tcB)

def cell_text(cell, text, size=10, bold=False, italic=False, center=False, color_rgb=None):
    par = cell.paragraphs[0]
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after  = Pt(3)
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)

# Color de cabecera: rojo oscuro igual al ejemplo
HDR_BG    = 'C00000'   # rojo oscuro
AREA_BG   = '1F3864'   # azul oscuro para celdas de área
ALT_BG    = 'F2F2F2'   # gris claro para filas alternas

# ══════════════════════════════════════════════════════════════════════════════
# TÍTULO
# ══════════════════════════════════════════════════════════════════════════════
h(doc, '6b. Resumen del Plan de Accion — Matriz de Control', level=2)

p(doc,
  'La siguiente matriz sintetiza las acciones correctivas que mitigan de forma '
  'etica y tecnica los hallazgos identificados durante la auditoria automatizada '
  'de Eco-Alerta. Esta presentacion esta diseñada para ser presentada ante una '
  'junta de auditoria o direccion del proyecto.',
  space_after=10)

# ══════════════════════════════════════════════════════════════════════════════
# MATRIZ PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
widths = [3.5, 5.5, 4.0, 4.0]
tm = doc.add_table(rows=1, cols=4)
tm.style = 'Table Grid'
tm.alignment = WD_TABLE_ALIGNMENT.CENTER

# Fila de encabezado
hdr_row = tm.rows[0]
for i, (txt, w) in enumerate(zip(
    ['Area de Mejora', 'Accion Automatizada Especifica',
     'Estandar / Norma de Referencia', 'Indicador de Exito (KPI)'],
    widths
)):
    cell = hdr_row.cells[i]
    cell.width = Cm(w)
    set_cell_bg(cell, HDR_BG)
    set_borders(cell, 'FFFFFF')
    cell_text(cell, txt, size=10, bold=True, center=True, color_rgb=(255,255,255))

# ── Datos de la matriz ────────────────────────────────────────────────────────
# Formato: (area, accion, estandar, kpi, bg_fila)
filas = [
    (
        'Seguridad',
        'Eliminar credenciales y tokens hardcodeados del codigo fuente. '
        'Implementar variables de entorno con python-decouple. '
        'Integrar git-secrets como hook de pre-commit para bloquear commits inseguros.',
        'ISO/IEC 25040\n(Caracteristica de Seguridad — Confidencialidad)\n\n'
        'OWASP Top 10 2021\n(A02: Cryptographic Failures)',
        '0 issues BLOCKER en SonarQube.\n'
        'Ningun archivo con credenciales en texto plano.\n'
        'Pre-commit hook activo y verificado.',
        'FFFFFF'
    ),
    (
        'Configuracion HTTP',
        'Configurar Content-Security-Policy, X-Frame-Options, HSTS y '
        'X-Content-Type-Options en Django. Asegurar atributos HttpOnly y '
        'SameSite en todas las cookies de sesion. Desactivar DEBUG=True '
        'en entorno de produccion.',
        'ISO/IEC 14764\n(Mantenimiento Perfectivo)\n\n'
        'OWASP Top 10 2021\n(A05: Security Misconfiguration)',
        'Re-escaneo OWASP ZAP sin alertas de nivel Medio.\n'
        'Cabeceras de seguridad presentes en 100% de las respuestas HTTP.',
        ALT_BG
    ),
    (
        'Control de Acceso',
        'Restringir metodos HTTP en endpoints de la API usando decoradores Django. '
        'Limitar CORS a dominios autorizados. Deshabilitar directory browsing. '
        'Eliminar tokens JWT de URLs y transmitirlos solo via Authorization header.',
        'ISO/IEC 25040\n(Caracteristica de Seguridad — Control de Acceso)\n\n'
        'OWASP Top 10 2021\n(A01: Broken Access Control)',
        '0 endpoints sin restriccion de metodo HTTP.\n'
        'CORS configurado solo para dominios autorizados.\n'
        'JWT transmitido exclusivamente via header.',
        'FFFFFF'
    ),
    (
        'Confiabilidad',
        'Refactorizar manejo de excepciones en el backend Python. '
        'Resolver los 44 issues High de Reliability identificados por SonarQube. '
        'Implementar suite de pruebas unitarias con pytest-django. '
        'Actualizar dependencias deprecadas del frontend React.',
        'ISO/IEC 25040\n(Caracteristica de Confiabilidad — Tolerancia a Fallos)\n\n'
        'ISO/IEC 12207\n(Proceso de Pruebas del Software)',
        'Issues High de Reliability reducidos en 80%.\n'
        'Cobertura de pruebas unitarias mayor o igual a 70% en modulos criticos.\n'
        'npm audit sin vulnerabilidades criticas.',
        ALT_BG
    ),
    (
        'Mantenibilidad',
        'Eliminar codigo duplicado aplicando el principio DRY. '
        'Documentar funciones criticas con docstrings. '
        'Eliminar imports y variables no usadas con flake8 y ESLint. '
        'Integrar SonarQube en pipeline CI/CD con Quality Gate bloqueante.',
        'ISO/IEC 25040\n(Caracteristica de Mantenibilidad — Modificabilidad)\n\n'
        'ISO/IEC 14764\n(Mantenimiento Perfectivo)',
        'SonarQube Quality Gate en estado Passed.\n'
        'Deuda tecnica reducida en 60%.\n'
        'Pipeline CI/CD rechaza PRs con nuevos issues BLOCKER o High.',
        'FFFFFF'
    ),
    (
        'Etica y Privacidad',
        'Revisar y eliminar comentarios con informacion sensible (TODO, FIXME, '
        'password, token) del codigo JavaScript del frontend. '
        'Asegurar que ningun dato PII de ciudadanos sea expuesto en logs, '
        'URLs ni mensajes de error del sistema.',
        'ISO/IEC 25040\n(Caracteristica de Seguridad — Privacidad)\n\n'
        'Ley N° 19.628\n(Proteccion de la Vida Privada — Chile)',
        '0 comentarios con informacion sensible en codigo de produccion.\n'
        'Logs del sistema sin datos PII de ciudadanos.\n'
        'Mensajes de error genericos en produccion (sin stack trace).',
        ALT_BG
    ),
]

for area, accion, estandar, kpi, bg in filas:
    row = tm.add_row()

    # Celda 0: Area (fondo azul oscuro, texto blanco, centrado)
    c0 = row.cells[0]
    c0.width = Cm(widths[0])
    set_cell_bg(c0, AREA_BG)
    set_borders(c0, 'FFFFFF')
    cell_text(c0, area, size=10, bold=True, center=True, color_rgb=(255,255,255))

    # Celda 1: Accion
    c1 = row.cells[1]
    c1.width = Cm(widths[1])
    set_cell_bg(c1, bg)
    set_borders(c1, 'CCCCCC')
    cell_text(c1, accion, size=9)

    # Celda 2: Estandar
    c2 = row.cells[2]
    c2.width = Cm(widths[2])
    set_cell_bg(c2, bg)
    set_borders(c2, 'CCCCCC')
    # Primero el nombre del estandar en negrita, resto normal
    par = c2.paragraphs[0]
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after  = Pt(3)
    for linea in estandar.split('\n'):
        if linea.startswith('ISO') or linea.startswith('OWASP') or linea.startswith('Ley'):
            run = par.add_run(linea + '\n')
            run.font.bold  = True
            run.font.size  = Pt(9)
            run.font.name  = 'Arial'
        elif linea.strip():
            run = par.add_run(linea + '\n')
            run.font.italic = True
            run.font.size   = Pt(8.5)
            run.font.name   = 'Arial'
        else:
            run = par.add_run('\n')
            run.font.size = Pt(4)
            run.font.name = 'Arial'

    # Celda 3: KPI
    c3 = row.cells[3]
    c3.width = Cm(widths[3])
    set_cell_bg(c3, bg)
    set_borders(c3, 'CCCCCC')
    par3 = c3.paragraphs[0]
    par3.paragraph_format.space_before = Pt(3)
    par3.paragraph_format.space_after  = Pt(3)
    for linea in kpi.split('\n'):
        if linea.strip():
            run = par3.add_run('• ' + linea + '\n')
            run.font.size = Pt(9)
            run.font.name = 'Arial'

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# NOTA FINAL
# ══════════════════════════════════════════════════════════════════════════════
p(doc,
  'Esta Matriz de Control debe ser revisada al finalizar cada sprint de '
  'implementacion, actualizando el estado de cada accion con base en los '
  'resultados de los re-escaneos de OWASP ZAP y SonarQube. El objetivo '
  'final es llevar todos los indicadores de exito a su meta definida antes '
  'del despliegue a produccion del sistema Eco-Alerta.',
  italic=True, size=10, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
out = '/sessions/tender-gifted-cray/mnt/ecoalerta/seccion6b_v2.docx'
doc.save(out)
print(f'OK: {out}')
