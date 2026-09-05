#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera seccion5_v1.docx — Sección 5: Documentación de Hallazgos"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3)
sec.right_margin  = Cm(2)

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)

for h_name, sz in [('Heading 1',14),('Heading 2',12),('Heading 3',11)]:
    s = doc.styles[h_name]
    s.font.name  = 'Arial'
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = RGBColor(0,0,0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'AAAAAA')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def table_header_row(table, headers, col_widths, bg='1F5C99'):
    row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size  = Pt(9)
        run.font.name  = 'Arial'

def add_data_row(table, values, col_widths, bg=None, sizes=None):
    row = table.add_row()
    for i, (val, w) in enumerate(zip(values, col_widths)):
        cell = row.cells[i]
        cell.width = Cm(w)
        if bg:
            set_cell_bg(cell, bg)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        sz = sizes[i] if sizes else 9
        run = p.add_run(str(val))
        run.font.size = Pt(sz)
        run.font.name = 'Arial'

def sev_color(sev):
    m = {
        'BLOCKER':'C00000','Blocker':'C00000',
        'MAJOR':'FF6600','Major':'FF6600','HIGH':'FF6600','High':'FF6600','CRITICAL':'FF6600',
        'MINOR':'70AD47','Minor':'70AD47','LOW':'70AD47','Low':'70AD47',
        'Medio':'FFC000','Medium':'FFC000','MEDIUM':'FFC000',
        'Bajo':'92D050','Informativo':'4472C4','Info':'4472C4',
    }
    return m.get(sev, 'AAAAAA')

# ══════════════════════════════════════════════════════════════════════════════
# TÍTULO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Documentación Obtenida al Ejecutar la Herramienta', level=1)

add_para(doc,
    'Esta sección registra de forma íntegra y estructurada los hallazgos, '
    'anomalías y vulnerabilidades detectadas tras la ejecución automatizada '
    'de las dos herramientas seleccionadas: OWASP ZAP 2.17.0 (DAST) y '
    'SonarQube Community Build 26.6.0 (SAST), sobre el sistema Eco-Alerta '
    'en su entorno local de desarrollo.', space_after=10)

# ── 5.1 Resumen ejecutivo ─────────────────────────────────────────────────────
add_heading(doc, '5.1 Resumen Ejecutivo de Resultados', level=2)
add_para(doc,
    'La ejecución combinada de ambas herramientas permitió identificar un '
    'total de 575 hallazgos distribuidos en los tres componentes del sistema '
    '(backend, frontend y código fuente), abarcando desde vulnerabilidades '
    'críticas de exposición de credenciales hasta deficiencias de configuración '
    'de cabeceras HTTP y problemas de mantenibilidad del código.', space_after=8)

t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t, ['Herramienta','Alcance','Total','Críticos/Altos','Fecha'],
                 [3.5, 3.5, 2.0, 3.0, 2.5])
for row_data in [
    ('OWASP ZAP 2.17.0',   'Backend  — http://localhost:8000',   '4',   '1 (Medio)',           '26/06/2026'),
    ('OWASP ZAP 2.17.0',   'Frontend — http://localhost:5173',   '18',  '4 (Medio)',            '26/06/2026'),
    ('SonarQube Community', 'Código fuente (backend + frontend)', '553', '3 Blocker + 44 High', '26/06/2026'),
]:
    add_data_row(t, row_data, [3.5,3.5,2.0,3.0,2.5], sizes=[9,9,9,9,9])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5.2 ZAP
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.2 Resultados OWASP ZAP — Análisis Dinámico (DAST)', level=2)
add_para(doc,
    'OWASP ZAP 2.17.0 fue ejecutado en modo escaneo activo sobre los dos '
    'puntos de entrada de Eco-Alerta. El escaneo del backend incluyó '
    'inyección de la cabecera de autorización JWT mediante la regla Replacer '
    'para simular sesiones autenticadas.', space_after=8)

# 5.2.1 Backend
add_heading(doc, '5.2.1 Escaneo Backend — http://localhost:8000', level=3)
add_para(doc,
    'Procesó 2 endpoints activos. El 85% de respuestas fueron 4xx '
    '(recursos protegidos). Se identificaron 4 alertas:', space_after=6)

tb = doc.add_table(rows=1, cols=6)
tb.style = 'Table Grid'
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tb,['ID','Alerta','Riesgo','CWE','URL Afectada','Evidencia / Impacto'],
                 [1.2,4.0,1.5,1.5,3.5,4.5], bg='1A4B8C')

for id_, alerta, riesgo, cwe, url, ev in [
    ('ZAP-B-01','Cabecera CSP no configurada','Medio','CWE-693',
     'GET /robots.txt',
     'Ausencia de Content-Security-Policy. Permite ataques XSS e inyección de contenido.'),
    ('ZAP-B-02','Cookie Sin Flag HttpOnly','Bajo','CWE-1004',
     'GET /robots.txt',
     'Set-Cookie: lab-session-id=12345-active; Path=/ (sin HttpOnly). Cookie accesible vía JS.'),
    ('ZAP-B-03','Cookie sin atributo SameSite','Bajo','CWE-1275',
     'GET /robots.txt',
     'Set-Cookie: lab-session-id (sin SameSite). Vulnerable a solicitudes CSRF cross-site.'),
    ('ZAP-B-04','Respuesta de Gestión de Sesión Identificada','Informativo','—',
     'GET /robots.txt',
     'cookie:lab-session-id detectado como token de gestión de sesión activo.'),
]:
    row = tb.add_row()
    for i, (val, w) in enumerate(zip([id_,alerta,riesgo,cwe,url,ev],
                                      [1.2,4.0,1.5,1.5,3.5,4.5])):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_borders(cell)
        if i == 2:
            set_cell_bg(cell, sev_color(riesgo))
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        run.font.name = 'Arial'
        if i == 2:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()
add_para(doc,
    'Nota: El servidor expone DEBUG = True en el mensaje de error 404 (Django), '
    'revelando la estructura interna de rutas del proyecto. Este hallazgo '
    'debe ser corregido antes de cualquier despliegue en producción.',
    italic=True, size=10, space_after=10)

# 5.2.2 Frontend
add_heading(doc, '5.2.2 Escaneo Frontend — http://localhost:5173', level=3)
add_para(doc,
    'El escaneo del frontend (React + Vite) detectó 18 alertas en '
    'http://localhost:5173. Se excluyen alertas de dominios externos de '
    'Microsoft Edge capturados por el proxy ZAP.', space_after=6)

# Tabla distribución
tf_dist = doc.add_table(rows=1, cols=4)
tf_dist.style = 'Table Grid'
tf_dist.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tf_dist, ['Riesgo','Cant.','% Total','Confianza Predominante'],
                 [3.0, 1.5, 2.0, 4.0], bg='1A4B8C')
for vals in [('Medio','4','22,2%','Alta / Media'),
             ('Bajo','8','44,4%','Alta / Media'),
             ('Informativo','6','33,3%','Media / Baja'),
             ('TOTAL','18','100%','—')]:
    add_data_row(tf_dist, vals, [3.0,1.5,2.0,4.0], sizes=[9,9,9,9])
doc.add_paragraph()

# Tabla detalle frontend
tf = doc.add_table(rows=1, cols=5)
tf.style = 'Table Grid'
tf.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tf, ['ID','Alerta','Riesgo','Ref.','Descripción y Recomendación'],
                 [1.2,4.0,1.5,1.8,7.5], bg='1A4B8C')

frontend_alerts = [
    ('ZAP-F-01','Cabecera CSP no configurada','Medio','CWE-693 / A05:2021',
     'Ausencia de Content-Security-Policy en 4 respuestas. Habilita ataques XSS. '
     'Solución: configurar CSP restrictiva en el servidor.'),
    ('ZAP-F-02','Configuración Incorrecta Cross-Domain','Medio','A05:2021',
     'Configuración CORS permisiva en localhost:5173. '
     'Solución: restringir Access-Control-Allow-Origin a dominios autorizados.'),
    ('ZAP-F-03','Directory Browsing','Medio','A05:2021',
     'Listado de directorio activo en /node_modules/cesium/. '
     'Solución: deshabilitar directory listing en servidor para producción.'),
    ('ZAP-F-04','Falta cabecera Anti-Clickjacking','Medio','A05:2021',
     'Ausencia de X-Frame-Options o frame-ancestors CSP en 4 respuestas. '
     'Solución: agregar X-Frame-Options: DENY.'),
    ('ZAP-F-05','Cookie Sin Flag HttpOnly','Bajo','CWE-1004',
     'Cookie lab-session-id sin HttpOnly en frontend. '
     'Solución: configurar HttpOnly en Django settings.'),
    ('ZAP-F-06','Cookie sin atributo SameSite','Bajo','CWE-1275',
     'Cookie lab-session-id sin SameSite. Vulnerable a CSRF. '
     'Solución: agregar SameSite=Strict o Lax.'),
    ('ZAP-F-07','Divulgación de Timestamps Unix','Bajo','CWE-200',
     'Timestamps Unix expuestos en 7 respuestas JS. Información de temporización filtrable.'),
    ('ZAP-F-08','Divulgación cabecera X-Powered-By','Bajo','CWE-200',
     'Header X-Powered-By revela tecnología del servidor. '
     'Solución: suprimir este header en configuración del servidor.'),
    ('ZAP-F-09','Falta cabecera X-Content-Type-Options','Bajo','A05:2021',
     'Ausencia de X-Content-Type-Options: nosniff en 12 respuestas. '
     'Permite MIME-sniffing. Solución: agregar header en todas las respuestas.'),
    ('ZAP-F-10','HSTS no configurado','Bajo','A05:2021',
     'Strict-Transport-Security ausente en 6 respuestas. '
     'En producción permite downgrade a HTTP. Activar HSTS con includeSubDomains.'),
    ('ZAP-F-11','Comentarios sospechosos en JS','Informativo','CWE-200',
     '15 instancias de comentarios con TODO, FIXME, token, password en código JS. '
     'Revisar y eliminar antes de producción.'),
    ('ZAP-F-12','Información sensible en URL','Informativo','CWE-200',
     'Parámetro token detectado en URL querystring. '
     'Los JWT no deben transmitirse por URL. Usar Authorization header.'),
]

for id_, alerta, riesgo, ref, desc in frontend_alerts:
    row = tf.add_row()
    for i, (val, w) in enumerate(zip([id_,alerta,riesgo,ref,desc],
                                      [1.2,4.0,1.5,1.8,7.5])):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_borders(cell)
        if i == 2:
            set_cell_bg(cell, sev_color(riesgo))
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        run.font.name = 'Arial'
        if i == 2:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5.3 SonarQube
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.3 Resultados SonarQube — Análisis Estático (SAST)', level=2)
add_para(doc,
    'SonarQube Community Build 26.6.0.123539 analizó 132 archivos fuente del '
    'proyecto Eco-Alerta (55 Python + 63 JavaScript/JSX + 1 Dockerfile + CSS). '
    'Análisis ejecutado el 26/06/2026 mediante sonar-scanner CLI 8.1.0.6389. '
    'Total: 553 issues detectados.', space_after=8)

# 5.3.1 Distribución
add_heading(doc, '5.3.1 Distribución General de Issues', level=3)

ts_dist = doc.add_table(rows=1, cols=5)
ts_dist.style = 'Table Grid'
ts_dist.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(ts_dist,
    ['Categoría','Total','Blocker','High/Critical','Medium/Low'],
    [4.0,1.5,1.5,2.5,2.5], bg='1A4B8C')
for vals in [
    ('Security (Seguridad)',        '9',   '3', '5',   '1'),
    ('Reliability (Confiabilidad)', '229', '0', '44',  '185'),
    ('Maintainability (Mantenimiento)','458','0','44*','414'),
    ('Info',                        '3',   '0', '0',   '3'),
    ('TOTAL',                       '553', '3', '44+', '280+'),
]:
    add_data_row(ts_dist, vals, [4.0,1.5,1.5,2.5,2.5], sizes=[9,9,9,9,9])
doc.add_paragraph()

add_para(doc,
    '* High en Maintainability corresponde a alta complejidad cognitiva y '
    'duplicación de código crítico.',
    italic=True, size=9, space_after=8)

# 5.3.2 Vulnerabilidades de seguridad
add_heading(doc, '5.3.2 Vulnerabilidades de Seguridad — Detalle (9 issues)', level=3)
add_para(doc,
    'A continuación se detallan los 9 issues catalogados como VULNERABILITY '
    'por SonarQube, clasificados por severidad:', space_after=6)

ts_sec = doc.add_table(rows=1, cols=6)
ts_sec.style = 'Table Grid'
ts_sec.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(ts_sec,
    ['ID','Archivo','Línea','Regla SonarQube','Severidad','Descripción del Hallazgo'],
    [1.2,4.5,1.0,2.0,1.8,5.5], bg='1A4B8C')

sq_vulns = [
    ('SQ-SEC-01','backend/reportes/views.py','2160','python:S6418','BLOCKER',
     'Token hardcodeado en código fuente (views.py). Posible exposición de credencial de API a cualquier persona con acceso al repositorio.'),
    ('SQ-SEC-02','backend/.../create_admin_user.py','31','python:S6437','BLOCKER',
     'Contraseña comprometida detectada en script de creación de admin. Debe ser revocada y eliminada del código de forma inmediata.'),
    ('SQ-SEC-03','backend/.../make_inspector_admin.py','26','python:S6437','BLOCKER',
     'Contraseña comprometida en script de gestión de roles. Requiere revocación urgente.'),
    ('SQ-SEC-04','backend/.../create_admin_user.py','31','python:S2068','MAJOR',
     'Credencial hardcodeada: campo "password" con valor literal en comando de gestión Django.'),
    ('SQ-SEC-05','backend/.../load_initial_data.py','44','python:S2068','MAJOR',
     'Contraseña hardcodeada en script de carga de datos iniciales (usuario inspector 1).'),
    ('SQ-SEC-06','backend/.../load_initial_data.py','73','python:S2068','MAJOR',
     'Contraseña hardcodeada en script de carga de datos iniciales (usuario inspector 2).'),
    ('SQ-SEC-07','backend/.../make_inspector_admin.py','26','python:S2068','MAJOR',
     'Contraseña hardcodeada en script de asignación de rol inspector/admin.'),
    ('SQ-SEC-08','backend/ecoalerta/urls.py','12','python:S3752','MINOR',
     'Endpoint Django no especifica métodos HTTP permitidos. Puede aceptar verbos no intencionados (PUT, DELETE).'),
    ('SQ-SEC-09','backend/ecoalerta/urls.py','30','python:S3752','MINOR',
     'Endpoint Django no especifica métodos HTTP permitidos. Misma vulnerabilidad en ruta distinta.'),
]

for id_, archivo, linea, regla, sev, desc in sq_vulns:
    row = ts_sec.add_row()
    for i, (val, w) in enumerate(zip([id_,archivo,linea,regla,sev,desc],
                                      [1.2,4.5,1.0,2.0,1.8,5.5])):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_borders(cell)
        if i == 4:
            set_cell_bg(cell, sev_color(sev))
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        run.font.name = 'Arial'
        if i == 4:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)

doc.add_paragraph()

# 5.3.3 Issues destacados
add_heading(doc, '5.3.3 Issues Destacados por Categoría', level=3)

ts_cat = doc.add_table(rows=1, cols=4)
ts_cat.style = 'Table Grid'
ts_cat.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(ts_cat,
    ['Categoría','Severidad','Cantidad','Patrón Detectado'],
    [3.0,2.0,1.5,9.5], bg='1A4B8C')
for vals in [
    ('Reliability','High','44',
     'Código que puede provocar comportamiento inesperado en runtime: manejo incorrecto de excepciones en Python, uso de variables potencialmente no inicializadas.'),
    ('Reliability','Medium/Low','185',
     'Uso de APIs deprecadas en React, condiciones de carrera potenciales y manejo inadecuado de errores en componentes frontend.'),
    ('Maintainability','High (complejidad)','~50',
     'Código duplicado: literal redis://127.0.0.1:6379/1 repetido 3 veces. Funciones con alta complejidad cognitiva. Clases con múltiples responsabilidades en views.py.'),
    ('Maintainability','Medium','~300',
     'Nomenclatura inconsistente, ausencia de docstrings en funciones críticas, imports no utilizados y variables sin usar en componentes JSX.'),
]:
    add_data_row(ts_cat, vals, [3.0,2.0,1.5,9.5], sizes=[9,9,9,9])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5.4 Consolidado
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.4 Consolidado de Hallazgos por Componente', level=2)
add_para(doc,
    'La siguiente tabla consolida todos los hallazgos por componente del '
    'sistema y los clasifica según OWASP Top 10 2021:', space_after=6)

tc = doc.add_table(rows=1, cols=5)
tc.style = 'Table Grid'
tc.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tc,
    ['Componente','Herramienta','Hallazgos','Clasificación OWASP 2021','Impacto Potencial'],
    [3.0,3.0,2.0,3.5,4.5])
for vals in [
    ('Backend — Django REST API',  'OWASP ZAP','4 alertas',
     'A05 Security Misconfiguration',
     'Exposición de configuración, riesgo de session hijacking, CSRF'),
    ('Backend — Código fuente',    'SonarQube', '7 vuln. (3 BLOCKER)',
     'A02 Cryptographic Failures + A05',
     'Credenciales hardcodeadas → acceso no autorizado al sistema en producción'),
    ('Backend — Deuda técnica',    'SonarQube', '229 Reliability\n458 Maint.',
     'A09 Security Logging Failures',
     'Alta deuda técnica, mantenibilidad comprometida, riesgo de fallos en runtime'),
    ('Frontend — React/Vite',      'OWASP ZAP', '18 alertas',
     'A05 Security Misconfiguration',
     'Clickjacking, CSRF, información sensible en URLs y comentarios JS'),
    ('API Endpoints — urls.py',    'SonarQube', '2 vuln.',
     'A01 Broken Access Control',
     'Endpoints sin restricción de método HTTP → posible acceso indebido'),
]:
    add_data_row(tc, vals, [3.0,3.0,2.0,3.5,4.5], sizes=[9,9,9,9,9])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5.5 Relación con auditoría manual
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.5 Relación con Hallazgos de Auditoría Manual (Eva #1)', level=2)
add_para(doc,
    'Las herramientas automatizadas complementan —pero no reemplazan— la '
    'revisión manual realizada en la Eva #1. Los 21 hallazgos identificados '
    'manualmente incluían vulnerabilidades de lógica de negocio (IDOR, '
    'escalación de privilegios, control de acceso por roles) que no son '
    'detectables por DAST/SAST. Sin embargo, la auditoría automatizada '
    'identificó adicionalmente vulnerabilidades de seguridad en código '
    'fuente (credenciales hardcodeadas, tokens expuestos) que no fueron '
    'detectadas en la revisión manual, demostrando el valor complementario '
    'de ambos enfoques.', space_after=8)

tm = doc.add_table(rows=1, cols=3)
tm.style = 'Table Grid'
tm.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tm,
    ['Tipo de Hallazgo','Detectado por','Cantidad'],
    [6.5, 4.5, 2.0])
for vals in [
    ('Vulnerabilidades de configuración HTTP (cabeceras)','ZAP — DAST','6'),
    ('Gestión insegura de cookies de sesión',             'ZAP — DAST','2'),
    ('Credenciales y tokens hardcodeados en código',      'SonarQube — SAST','7'),
    ('Métodos HTTP no restringidos en endpoints',         'SonarQube — SAST','2'),
    ('Deficiencias de confiabilidad y mantenibilidad',    'SonarQube — SAST','549'),
    ('Vulnerabilidades de lógica de negocio (IDOR, etc.)','Revisión manual — Eva #1','21'),
]:
    add_data_row(tm, vals, [6.5,4.5,2.0], sizes=[9,9,9])

doc.add_paragraph()

add_para(doc,
    'Todos los hallazgos documentados en esta sección serán abordados en la '
    'Sección 6 mediante un Plan de Mejoras alineado a los estándares '
    'ISO/IEC 25040, ISO/IEC 12207 e ISO/IEC 14764.',
    italic=True, size=10, space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
out = '/sessions/tender-gifted-cray/mnt/ecoalerta/seccion5_v1.docx'
doc.save(out)
print(f'OK: {out}')
